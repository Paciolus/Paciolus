"""
Tests for shared.docx_parser — DOCX File Parsing & ZIP Disambiguation.

Tests cover:
- DOCX parsing (valid file, multi-table, metadata attachment)
- ZIP disambiguation (_is_docx_zip for DOCX vs XLSX vs ODS)
- Error handling (corrupt, no tables, single-row table)
- Integration with parse_uploaded_file_by_format()
- Format detection integration
"""

import io
import zipfile

import pandas as pd
import pytest
from fastapi import HTTPException

from shared.docx_parser import DocxMetadata, _is_docx_zip, parse_docx

# =============================================================================
# Helpers — Build minimal DOCX and XLSX ZIPs in memory
# =============================================================================


def _build_docx_with_table(
    headers: list[str] | None = None,
    rows: list[list[str]] | None = None,
) -> bytes:
    """Build a minimal DOCX file with one table using python-docx."""
    from docx import Document

    if headers is None:
        headers = ["Account", "Debit", "Credit"]
    if rows is None:
        rows = [
            ["1000 - Cash", "50000", "0"],
            ["2000 - AP", "0", "30000"],
            ["3000 - Equity", "0", "20000"],
        ]

    doc = Document()
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))

    # Header row
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, val in enumerate(row_data):
            table.rows[row_idx + 1].cells[col_idx].text = val

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_docx_no_tables() -> bytes:
    """Build a DOCX file with only paragraphs, no tables."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("This document has no tables.")
    doc.add_paragraph("Just some text content.")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_docx_multi_table() -> bytes:
    """Build a DOCX file with multiple tables."""
    from docx import Document

    doc = Document()

    # First table: small header-only table
    t1 = doc.add_table(rows=1, cols=2)
    t1.rows[0].cells[0].text = "Title"
    t1.rows[0].cells[1].text = "Value"

    doc.add_paragraph("Some text between tables.")

    # Second table: actual data table
    t2 = doc.add_table(rows=3, cols=3)
    t2.rows[0].cells[0].text = "Account"
    t2.rows[0].cells[1].text = "Debit"
    t2.rows[0].cells[2].text = "Credit"
    t2.rows[1].cells[0].text = "1000"
    t2.rows[1].cells[1].text = "100"
    t2.rows[1].cells[2].text = "0"
    t2.rows[2].cells[0].text = "2000"
    t2.rows[2].cells[1].text = "0"
    t2.rows[2].cells[2].text = "100"

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_xlsx_zip() -> bytes:
    """Build a minimal XLSX file as bytes."""
    buf = io.BytesIO()
    df = pd.DataFrame({"Col1": [1, 2], "Col2": [3, 4]})
    df.to_excel(buf, index=False, engine="openpyxl")
    return buf.getvalue()


def _build_raw_zip_with_word_dir() -> bytes:
    """Build a raw ZIP with word/ directory (DOCX-like)."""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("word/document.xml", "<root/>")
        zf.writestr("[Content_Types].xml", "<root/>")
    return buf.getvalue()


def _build_raw_zip_with_xl_dir() -> bytes:
    """Build a raw ZIP with xl/ directory (XLSX-like)."""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("xl/workbook.xml", "<root/>")
        zf.writestr("[Content_Types].xml", "<root/>")
    return buf.getvalue()


# =============================================================================
# _is_docx_zip() — ZIP disambiguation
# =============================================================================


class TestIsDocxZip:
    """ZIP content inspection for DOCX vs XLSX disambiguation."""

    def test_docx_with_word_dir(self):
        docx_bytes = _build_raw_zip_with_word_dir()
        assert _is_docx_zip(docx_bytes) is True

    def test_xlsx_with_xl_dir(self):
        xlsx_bytes = _build_raw_zip_with_xl_dir()
        assert _is_docx_zip(xlsx_bytes) is False

    def test_real_docx_file(self):
        docx_bytes = _build_docx_with_table()
        assert _is_docx_zip(docx_bytes) is True

    def test_real_xlsx_file(self):
        xlsx_bytes = _build_xlsx_zip()
        assert _is_docx_zip(xlsx_bytes) is False

    def test_invalid_zip_returns_false(self):
        assert _is_docx_zip(b"not a zip file at all") is False

    def test_empty_zip(self):
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w") as zf:
            pass  # empty ZIP
        assert _is_docx_zip(buf.getvalue()) is False

    def test_zip_with_both_word_and_xl_is_not_docx(self):
        """ZIP with both word/ and xl/ directories — xl/ takes precedence."""
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w") as zf:
            zf.writestr("word/document.xml", "<root/>")
            zf.writestr("xl/workbook.xml", "<root/>")
        assert _is_docx_zip(buf.getvalue()) is False


# =============================================================================
# parse_docx() — Core parsing
# =============================================================================


class TestParseDocx:
    """DOCX file parsing via python-docx."""

    def test_parse_valid_single_table(self):
        docx_bytes = _build_docx_with_table()
        df = parse_docx(docx_bytes, "test.docx")
        assert len(df) == 3
        assert "Account" in df.columns
        assert "Debit" in df.columns
        assert "Credit" in df.columns

    def test_data_values_correct(self):
        docx_bytes = _build_docx_with_table()
        df = parse_docx(docx_bytes, "test.docx")
        assert df.iloc[0]["Account"] == "1000 - Cash"
        assert df.iloc[0]["Debit"] == "50000"
        assert df.iloc[1]["Credit"] == "30000"

    def test_metadata_attached(self):
        docx_bytes = _build_docx_with_table()
        df = parse_docx(docx_bytes, "test.docx")
        metadata = df.attrs.get("docx_metadata")
        assert metadata is not None
        assert isinstance(metadata, DocxMetadata)
        assert metadata.table_count == 1
        assert metadata.table_index == 0
        assert metadata.original_row_count == 3
        assert metadata.original_col_count == 3

    def test_metadata_is_frozen(self):
        docx_bytes = _build_docx_with_table()
        df = parse_docx(docx_bytes, "test.docx")
        metadata = df.attrs["docx_metadata"]
        with pytest.raises(AttributeError):
            metadata.table_count = 99

    def test_no_tables_raises_400(self):
        docx_bytes = _build_docx_no_tables()
        with pytest.raises(HTTPException) as exc_info:
            parse_docx(docx_bytes, "no_tables.docx")
        assert exc_info.value.status_code == 400
        assert "no tables" in exc_info.value.detail

    def test_multi_table_skips_header_only(self):
        """First table has only headers (1 row), parser should pick second table."""
        docx_bytes = _build_docx_multi_table()
        df = parse_docx(docx_bytes, "multi.docx")
        # Second table has Account/Debit/Credit
        assert "Account" in df.columns
        assert len(df) == 2
        metadata = df.attrs["docx_metadata"]
        assert metadata.table_count == 2
        assert metadata.table_index == 1  # 0-based, second table

    def test_corrupt_file_raises_400(self):
        with pytest.raises(HTTPException) as exc_info:
            parse_docx(b"not a docx file", "corrupt.docx")
        assert exc_info.value.status_code == 400
        assert "could not be opened" in exc_info.value.detail

    def test_returns_dataframe(self):
        docx_bytes = _build_docx_with_table()
        result = parse_docx(docx_bytes, "test.docx")
        assert isinstance(result, pd.DataFrame)

    def test_empty_headers_get_placeholder(self):
        """Columns with empty header text should get 'Column N' placeholders."""
        docx_bytes = _build_docx_with_table(
            headers=["Account", "", "Balance"],
            rows=[["1000", "", "500"]],
        )
        df = parse_docx(docx_bytes, "test.docx")
        assert "Column 2" in df.columns

    def test_single_row_table_returns_empty_df(self):
        """Table with only headers (no data rows) should return empty DataFrame."""
        docx_bytes = _build_docx_with_table(headers=["A", "B"], rows=[])
        df = parse_docx(docx_bytes, "single_row.docx")
        assert len(df) == 0
        assert list(df.columns) == ["A", "B"]


# =============================================================================
# Format detection integration
# =============================================================================


class TestDocxFormatDetection:
    """DOCX detection via detect_format() and file_formats module."""

    def test_detect_by_extension(self):
        from shared.file_formats import FileFormat, detect_format

        result = detect_format(filename="data.docx")
        assert result.format == FileFormat.DOCX
        assert result.confidence == "high"
        assert result.source == "extension"

    def test_detect_docx_by_magic_bytes(self):
        from shared.file_formats import FileFormat, detect_format

        docx_bytes = _build_docx_with_table()
        result = detect_format(file_bytes=docx_bytes)
        assert result.format == FileFormat.DOCX
        assert result.source == "magic"

    def test_detect_xlsx_by_magic_bytes_not_docx(self):
        from shared.file_formats import FileFormat, detect_format

        xlsx_bytes = _build_xlsx_zip()
        result = detect_format(file_bytes=xlsx_bytes)
        assert result.format == FileFormat.XLSX
        assert result.source == "magic"

    def test_extension_overrides_zip_inspection(self):
        """Extension has higher priority than magic bytes."""
        from shared.file_formats import FileFormat, detect_format

        xlsx_bytes = _build_xlsx_zip()
        result = detect_format(filename="report.docx", file_bytes=xlsx_bytes)
        assert result.format == FileFormat.DOCX
        assert result.source == "extension"

    def test_docx_content_type_detection(self):
        from shared.file_formats import FileFormat, detect_format

        result = detect_format(content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        assert result.format == FileFormat.DOCX
        assert result.source == "content_type"


# =============================================================================
# Profile assertions
# =============================================================================


class TestDocxProfile:
    """DOCX profile in FORMAT_PROFILES."""

    def test_docx_parse_supported(self):
        from shared.file_formats import FORMAT_PROFILES, FileFormat

        profile = FORMAT_PROFILES[FileFormat.DOCX]
        assert profile.parse_supported is True

    def test_docx_content_types(self):
        from shared.file_formats import FORMAT_PROFILES, FileFormat

        profile = FORMAT_PROFILES[FileFormat.DOCX]
        assert "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in profile.content_types
        assert "application/octet-stream" in profile.content_types

    def test_docx_extension(self):
        from shared.file_formats import FORMAT_PROFILES, FileFormat

        profile = FORMAT_PROFILES[FileFormat.DOCX]
        assert ".docx" in profile.extensions

    def test_docx_label(self):
        from shared.file_formats import FORMAT_PROFILES, FileFormat

        profile = FORMAT_PROFILES[FileFormat.DOCX]
        assert profile.label == "Word (.docx)"

    def test_docx_in_allowed_extensions(self):
        from shared.file_formats import ALLOWED_EXTENSIONS

        assert ".docx" in ALLOWED_EXTENSIONS

    def test_docx_in_allowed_content_types(self):
        from shared.file_formats import ALLOWED_CONTENT_TYPES

        assert "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in ALLOWED_CONTENT_TYPES

    def test_display_includes_docx(self):
        from shared.file_formats import get_active_extensions_display

        display = get_active_extensions_display()
        assert ".docx" in display
